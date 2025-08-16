"""
Report Generator
===============

Generates comprehensive automation reports in Word, Excel, and PDF formats
with Playwright code, screenshots, execution details, and analysis.
"""

import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path
import base64
from io import BytesIO

# Report generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Create mock Document class
    class Document:
        def __init__(self, *args, **kwargs): pass
        def add_table(self, *args, **kwargs): 
            class MockTable:
                style = None
                def add_row(self): 
                    class MockRow:
                        cells = [type('MockCell', (), {'text': '', 'paragraphs': []})() for _ in range(2)]
                    return MockRow()
            return MockTable()
        def add_paragraph(self, *args, **kwargs): 
            class MockParagraph:
                def add_run(self, text): 
                    class MockRun:
                        bold = False
                        font = type('MockFont', (), {'size': None})()
                    return MockRun()
            return MockParagraph()
        def save(self, *args, **kwargs): pass
    
    class Inches:
        def __init__(self, value): pass
    
    class Pt:
        def __init__(self, value): pass

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

EXCEL_AVAILABLE = PANDAS_AVAILABLE and OPENPYXL_AVAILABLE

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class ReportGenerator:
    """
    Comprehensive report generator for automation results.
    """
    
    def __init__(self, config: Any):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.reports_dir = Path(config.database.media_path) / "reports"
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        
        # Check library availability
        self.libraries_available = {
            "docx": DOCX_AVAILABLE,
            "excel": EXCEL_AVAILABLE,
            "pdf": PDF_AVAILABLE
        }
        
        self.logger.info(f"Report generator initialized. Libraries: {self.libraries_available}")
    
    async def generate_comprehensive_report(
        self, 
        automation_result: Dict[str, Any], 
        formats: List[str] = ["docx", "excel", "pdf"]
    ) -> Dict[str, str]:
        """
        Generate comprehensive automation report in multiple formats.
        
        Args:
            automation_result: Complete automation execution result
            formats: List of output formats ("docx", "excel", "pdf")
            
        Returns:
            Dictionary with file paths for each generated format
        """
        try:
            self.logger.info(f"Generating comprehensive report in formats: {formats}")
            
            # Prepare report data
            report_data = self._prepare_report_data(automation_result)
            
            # Generate reports
            generated_files = {}
            
            if "docx" in formats and self.libraries_available["docx"]:
                docx_path = await self._generate_word_report(report_data)
                generated_files["docx"] = docx_path
                
            if "excel" in formats and self.libraries_available["excel"]:
                excel_path = await self._generate_excel_report(report_data)
                generated_files["excel"] = excel_path
                
            if "pdf" in formats and self.libraries_available["pdf"]:
                pdf_path = await self._generate_pdf_report(report_data)
                generated_files["pdf"] = pdf_path
            
            self.logger.info(f"Generated reports: {list(generated_files.keys())}")
            return generated_files
            
        except Exception as e:
            self.logger.error(f"Failed to generate comprehensive report: {e}", exc_info=True)
            return {}
    
    def _prepare_report_data(self, automation_result: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare and structure data for report generation."""
        try:
            # Extract key information
            instructions = automation_result.get("instructions", "")
            url = automation_result.get("url", "")
            status = automation_result.get("status", "unknown")
            timestamp = automation_result.get("timestamp", datetime.utcnow().isoformat())
            
            # Parse timestamp
            try:
                dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                formatted_time = dt.strftime("%Y-%m-%d %H:%M:%S UTC")
            except:
                formatted_time = timestamp
            
            # Extract execution details
            execution_result = automation_result.get("results", {})
            automation_plan = automation_result.get("automation_plan", {})
            ai_dom_analysis = automation_result.get("ai_dom_analysis", {})
            
            # Prepare screenshots
            screenshots = automation_result.get("screenshots", [])
            if isinstance(screenshots, dict):
                screenshots = list(screenshots.values())
            
            # Generate Playwright code
            playwright_code = self._generate_playwright_code(automation_result)
            
            # Prepare report data structure
            report_data = {
                "title": f"Automation Report - {instructions[:50]}...",
                "instructions": instructions,
                "url": url,
                "status": status,
                "timestamp": formatted_time,
                "execution_time": execution_result.get("execution_time", 0),
                "steps": execution_result.get("steps", []),
                "automation_plan": automation_plan,
                "ai_analysis": ai_dom_analysis,
                "screenshots": screenshots,
                "playwright_code": playwright_code,
                "error_details": automation_result.get("error", ""),
                "success_rate": self._calculate_success_rate(execution_result),
                "performance_metrics": self._extract_performance_metrics(execution_result),
                "ai_insights": self._extract_ai_insights(ai_dom_analysis)
            }
            
            return report_data
            
        except Exception as e:
            self.logger.error(f"Failed to prepare report data: {e}")
            return {}
    
    def _generate_playwright_code(self, automation_result: Dict[str, Any]) -> str:
        """Generate complete Playwright code from automation result."""
        try:
            instructions = automation_result.get("instructions", "")
            url = automation_result.get("url", "")
            steps = automation_result.get("results", {}).get("steps", [])
            
            code_lines = [
                "// Generated Playwright Code",
                "// Automation: " + instructions,
                "// URL: " + url,
                "// Generated: " + datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC"),
                "",
                "import { test, expect } from '@playwright/test';",
                "",
                "test('Automated Test', async ({ page }) => {",
                "  // Navigate to the target URL",
                f"  await page.goto('{url}');",
                "",
                "  // Wait for page to load",
                "  await page.waitForLoadState('networkidle');",
                ""
            ]
            
            # Add steps
            for i, step in enumerate(steps, 1):
                action = step.get("action", "")
                selector = step.get("selector", "")
                value = step.get("value", "")
                description = step.get("description", "")
                
                code_lines.append(f"  // Step {i}: {description}")
                
                if action == "click":
                    code_lines.append(f"  await page.click('{selector}');")
                elif action == "type":
                    code_lines.append(f"  await page.fill('{selector}', '{value}');")
                elif action == "wait":
                    duration = step.get("duration", 1)
                    code_lines.append(f"  await page.waitForTimeout({duration * 1000});")
                elif action == "navigate":
                    code_lines.append(f"  await page.goto('{value}');")
                else:
                    code_lines.append(f"  // {action} action with selector: {selector}")
                
                code_lines.append("")
            
            code_lines.extend([
                "  // Take final screenshot",
                "  await page.screenshot({ path: 'final-screenshot.png' });",
                "",
                "  // Verify completion",
                "  console.log('Automation completed successfully');",
                "});"
            ])
            
            return "\n".join(code_lines)
            
        except Exception as e:
            self.logger.error(f"Failed to generate Playwright code: {e}")
            return "// Error generating Playwright code"
    
    def _calculate_success_rate(self, execution_result: Dict[str, Any]) -> float:
        """Calculate success rate from execution results."""
        try:
            steps = execution_result.get("steps", [])
            if not steps:
                return 0.0
            
            successful_steps = sum(1 for step in steps if step.get("status") == "success")
            return (successful_steps / len(steps)) * 100
            
        except Exception as e:
            self.logger.error(f"Failed to calculate success rate: {e}")
            return 0.0
    
    def _extract_performance_metrics(self, execution_result: Dict[str, Any]) -> Dict[str, Any]:
        """Extract performance metrics from execution results."""
        try:
            steps = execution_result.get("steps", [])
            
            metrics = {
                "total_steps": len(steps),
                "execution_time": execution_result.get("execution_time", 0),
                "average_step_time": 0,
                "fastest_step": 0,
                "slowest_step": 0,
                "successful_steps": 0,
                "failed_steps": 0
            }
            
            if steps:
                step_times = [step.get("duration", 0) for step in steps]
                metrics["average_step_time"] = sum(step_times) / len(step_times) if step_times else 0
                metrics["fastest_step"] = min(step_times) if step_times else 0
                metrics["slowest_step"] = max(step_times) if step_times else 0
                metrics["successful_steps"] = sum(1 for step in steps if step.get("status") == "success")
                metrics["failed_steps"] = sum(1 for step in steps if step.get("status") == "failed")
            
            return metrics
            
        except Exception as e:
            self.logger.error(f"Failed to extract performance metrics: {e}")
            return {}
    
    def _extract_ai_insights(self, ai_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Extract AI insights from DOM analysis."""
        try:
            insights = {
                "elements_analyzed": 0,
                "login_elements_found": 0,
                "form_elements_found": 0,
                "navigation_elements_found": 0,
                "interactive_elements_found": 0,
                "confidence_score": 0.0,
                "recommendations": []
            }
            
            # Extract element counts
            insights["login_elements_found"] = len(ai_analysis.get("login_elements", []))
            insights["form_elements_found"] = len(ai_analysis.get("form_elements", []))
            insights["navigation_elements_found"] = len(ai_analysis.get("navigation_elements", []))
            insights["interactive_elements_found"] = len(ai_analysis.get("interactive_elements", []))
            
            insights["elements_analyzed"] = (
                insights["login_elements_found"] + 
                insights["form_elements_found"] + 
                insights["navigation_elements_found"] + 
                insights["interactive_elements_found"]
            )
            
            # Extract confidence score
            ai_analysis_data = ai_analysis.get("ai_analysis", {})
            insights["confidence_score"] = ai_analysis_data.get("confidence_score", 0.0)
            
            # Extract recommendations
            recommendations = ai_analysis_data.get("recommended_actions", [])
            insights["recommendations"] = [
                rec.get("description", "") for rec in recommendations
            ]
            
            return insights
            
        except Exception as e:
            self.logger.error(f"Failed to extract AI insights: {e}")
            return {}
    
    async def _generate_word_report(self, report_data: Dict[str, Any]) -> str:
        """Generate comprehensive Word document report."""
        try:
            if not self.libraries_available["docx"]:
                raise ImportError("python-docx library not available")
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(report_data["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata table
            self._add_metadata_table(doc, report_data)
            
            # Add executive summary
            self._add_executive_summary(doc, report_data)
            
            # Add execution details
            self._add_execution_details(doc, report_data)
            
            # Add AI analysis
            self._add_ai_analysis(doc, report_data)
            
            # Add screenshots
            self._add_screenshots(doc, report_data)
            
            # Add Playwright code
            self._add_playwright_code(doc, report_data)
            
            # Add performance metrics
            self._add_performance_metrics(doc, report_data)
            
            # Save document
            filename = f"automation_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.reports_dir / filename
            doc.save(str(filepath))
            
            self.logger.info(f"Word report generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Failed to generate Word report: {e}", exc_info=True)
            raise
    
    def _add_metadata_table(self, doc: Document, report_data: Dict[str, Any]):
        """Add metadata table to Word document."""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Property"
        header_cells[1].text = "Value"
        
        # Add data
        data = [
            ("Instructions", report_data["instructions"]),
            ("URL", report_data["url"]),
            ("Status", report_data["status"]),
            ("Timestamp", report_data["timestamp"]),
            ("Execution Time", f"{report_data['execution_time']:.2f} seconds"),
            ("Success Rate", f"{report_data['success_rate']:.1f}%")
        ]
        
        for property_name, value in data:
            row_cells = table.add_row().cells
            row_cells[0].text = property_name
            row_cells[1].text = str(value)
    
    def _add_executive_summary(self, doc: Document, report_data: Dict[str, Any]):
        """Add executive summary to Word document."""
        doc.add_heading("Executive Summary", level=1)
        
        summary = doc.add_paragraph()
        summary.add_run("Automation Status: ").bold = True
        summary.add_run(report_data["status"].upper())
        
        if report_data["status"] == "completed":
            summary.add_run("\n\n✅ The automation completed successfully with ")
            summary.add_run(f"{report_data['success_rate']:.1f}% success rate")
            summary.add_run(f" in {report_data['execution_time']:.2f} seconds.")
        else:
            summary.add_run("\n\n❌ The automation encountered issues during execution.")
            if report_data["error_details"]:
                summary.add_run(f"\nError: {report_data['error_details']}")
        
        # Add AI insights
        insights = report_data["ai_insights"]
        if insights["elements_analyzed"] > 0:
            summary.add_run(f"\n\n🤖 AI Analysis: Analyzed {insights['elements_analyzed']} elements ")
            summary.add_run(f"with {insights['confidence_score']:.1f}% confidence.")
    
    def _add_execution_details(self, doc: Document, report_data: Dict[str, Any]):
        """Add execution details to Word document."""
        doc.add_heading("Execution Details", level=1)
        
        # Add steps table
        if report_data["steps"]:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = "Step"
            header_cells[1].text = "Action"
            header_cells[2].text = "Description"
            header_cells[3].text = "Status"
            
            # Add steps
            for i, step in enumerate(report_data["steps"], 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = step.get("action", "")
                row_cells[2].text = step.get("description", "")
                row_cells[3].text = step.get("status", "")
    
    def _add_ai_analysis(self, doc: Document, report_data: Dict[str, Any]):
        """Add AI analysis to Word document."""
        doc.add_heading("AI Analysis", level=1)
        
        insights = report_data["ai_insights"]
        
        # Add element analysis
        p = doc.add_paragraph()
        p.add_run("Element Analysis:\n").bold = True
        p.add_run(f"• Login Elements: {insights['login_elements_found']}\n")
        p.add_run(f"• Form Elements: {insights['form_elements_found']}\n")
        p.add_run(f"• Navigation Elements: {insights['navigation_elements_found']}\n")
        p.add_run(f"• Interactive Elements: {insights['interactive_elements_found']}\n")
        p.add_run(f"• Total Elements Analyzed: {insights['elements_analyzed']}\n")
        p.add_run(f"• AI Confidence Score: {insights['confidence_score']:.1f}%\n")
        
        # Add recommendations
        if insights["recommendations"]:
            p = doc.add_paragraph()
            p.add_run("AI Recommendations:\n").bold = True
            for rec in insights["recommendations"]:
                p.add_run(f"• {rec}\n")
    
    def _add_screenshots(self, doc: Document, report_data: Dict[str, Any]):
        """Add screenshots to Word document."""
        doc.add_heading("Screenshots", level=1)
        
        screenshots = report_data["screenshots"]
        if not screenshots:
            doc.add_paragraph("No screenshots available.")
            return
        
        for i, screenshot_path in enumerate(screenshots, 1):
            try:
                if os.path.exists(screenshot_path):
                    doc.add_paragraph(f"Screenshot {i}:")
                    doc.add_picture(screenshot_path, width=Inches(6))
                    doc.add_paragraph()
            except Exception as e:
                self.logger.warning(f"Failed to add screenshot {screenshot_path}: {e}")
    
    def _add_playwright_code(self, doc: Document, report_data: Dict[str, Any]):
        """Add Playwright code to Word document."""
        doc.add_heading("Generated Playwright Code", level=1)
        
        # Add code as preformatted text
        code_paragraph = doc.add_paragraph()
        code_run = code_paragraph.add_run(report_data["playwright_code"])
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)
    
    def _add_performance_metrics(self, doc: Document, report_data: Dict[str, Any]):
        """Add performance metrics to Word document."""
        doc.add_heading("Performance Metrics", level=1)
        
        metrics = report_data["performance_metrics"]
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Value"
        
        # Add metrics
        metric_data = [
            ("Total Steps", str(metrics["total_steps"])),
            ("Successful Steps", str(metrics["successful_steps"])),
            ("Failed Steps", str(metrics["failed_steps"])),
            ("Execution Time", f"{metrics['execution_time']:.2f} seconds"),
            ("Average Step Time", f"{metrics['average_step_time']:.2f} seconds"),
            ("Fastest Step", f"{metrics['fastest_step']:.2f} seconds"),
            ("Slowest Step", f"{metrics['slowest_step']:.2f} seconds")
        ]
        
        for metric_name, value in metric_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric_name
            row_cells[1].text = value
    
    async def _generate_excel_report(self, report_data: Dict[str, Any]) -> str:
        """Generate comprehensive Excel report."""
        try:
            if not self.libraries_available["excel"]:
                raise ImportError("pandas/openpyxl libraries not available")
            
            # Create Excel workbook
            wb = openpyxl.Workbook()
            
            # Remove default sheet
            wb.remove(wb.active)
            
            # Add summary sheet
            self._add_excel_summary_sheet(wb, report_data)
            
            # Add execution details sheet
            self._add_excel_execution_sheet(wb, report_data)
            
            # Add AI analysis sheet
            self._add_excel_ai_sheet(wb, report_data)
            
            # Add performance metrics sheet
            self._add_excel_performance_sheet(wb, report_data)
            
            # Add Playwright code sheet
            self._add_excel_code_sheet(wb, report_data)
            
            # Save workbook
            filename = f"automation_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = self.reports_dir / filename
            wb.save(str(filepath))
            
            self.logger.info(f"Excel report generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Failed to generate Excel report: {e}", exc_info=True)
            raise
    
    def _add_excel_summary_sheet(self, wb: openpyxl.Workbook, report_data: Dict[str, Any]):
        """Add summary sheet to Excel workbook."""
        ws = wb.create_sheet("Summary")
        
        # Add title
        ws['A1'] = "Automation Report Summary"
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:C1')
        
        # Add metadata
        data = [
            ("Instructions", report_data["instructions"]),
            ("URL", report_data["url"]),
            ("Status", report_data["status"]),
            ("Timestamp", report_data["timestamp"]),
            ("Execution Time", f"{report_data['execution_time']:.2f} seconds"),
            ("Success Rate", f"{report_data['success_rate']:.1f}%")
        ]
        
        for i, (key, value) in enumerate(data, 3):
            ws[f'A{i}'] = key
            ws[f'B{i}'] = value
            ws[f'A{i}'].font = Font(bold=True)
    
    def _add_excel_execution_sheet(self, wb: openpyxl.Workbook, report_data: Dict[str, Any]):
        """Add execution details sheet to Excel workbook."""
        ws = wb.create_sheet("Execution Details")
        
        # Add headers
        headers = ["Step", "Action", "Description", "Status", "Duration", "Selector"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        
        # Add steps
        for i, step in enumerate(report_data["steps"], 2):
            ws.cell(row=i, column=1, value=i-1)
            ws.cell(row=i, column=2, value=step.get("action", ""))
            ws.cell(row=i, column=3, value=step.get("description", ""))
            ws.cell(row=i, column=4, value=step.get("status", ""))
            ws.cell(row=i, column=5, value=step.get("duration", 0))
            ws.cell(row=i, column=6, value=step.get("selector", ""))
    
    def _add_excel_ai_sheet(self, wb: openpyxl.Workbook, report_data: Dict[str, Any]):
        """Add AI analysis sheet to Excel workbook."""
        ws = wb.create_sheet("AI Analysis")
        
        insights = report_data["ai_insights"]
        
        # Add element analysis
        ws['A1'] = "Element Analysis"
        ws['A1'].font = Font(size=14, bold=True)
        
        element_data = [
            ("Login Elements", insights["login_elements_found"]),
            ("Form Elements", insights["form_elements_found"]),
            ("Navigation Elements", insights["navigation_elements_found"]),
            ("Interactive Elements", insights["interactive_elements_found"]),
            ("Total Elements Analyzed", insights["elements_analyzed"]),
            ("AI Confidence Score", f"{insights['confidence_score']:.1f}%")
        ]
        
        for i, (element_type, count) in enumerate(element_data, 3):
            ws[f'A{i}'] = element_type
            ws[f'B{i}'] = count
            ws[f'A{i}'].font = Font(bold=True)
    
    def _add_excel_performance_sheet(self, wb: openpyxl.Workbook, report_data: Dict[str, Any]):
        """Add performance metrics sheet to Excel workbook."""
        ws = wb.create_sheet("Performance Metrics")
        
        metrics = report_data["performance_metrics"]
        
        # Add metrics
        metric_data = [
            ("Total Steps", metrics["total_steps"]),
            ("Successful Steps", metrics["successful_steps"]),
            ("Failed Steps", metrics["failed_steps"]),
            ("Execution Time (seconds)", metrics["execution_time"]),
            ("Average Step Time (seconds)", metrics["average_step_time"]),
            ("Fastest Step (seconds)", metrics["fastest_step"]),
            ("Slowest Step (seconds)", metrics["slowest_step"])
        ]
        
        for i, (metric, value) in enumerate(metric_data, 1):
            ws[f'A{i}'] = metric
            ws[f'B{i}'] = value
            ws[f'A{i}'].font = Font(bold=True)
    
    def _add_excel_code_sheet(self, wb: openpyxl.Workbook, report_data: Dict[str, Any]):
        """Add Playwright code sheet to Excel workbook."""
        ws = wb.create_sheet("Playwright Code")
        
        # Add code
        ws['A1'] = "Generated Playwright Code"
        ws['A1'].font = Font(size=14, bold=True)
        
        # Split code into lines and add to cells
        code_lines = report_data["playwright_code"].split('\n')
        for i, line in enumerate(code_lines, 3):
            ws[f'A{i}'] = line
            ws[f'A{i}'].font = Font(name='Courier New', size=9)
    
    async def _generate_pdf_report(self, report_data: Dict[str, Any]) -> str:
        """Generate comprehensive PDF report."""
        try:
            if not self.libraries_available["pdf"]:
                raise ImportError("reportlab library not available")
            
            # Create PDF document
            filename = f"automation_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
            filepath = self.reports_dir / filename
            doc = SimpleDocTemplate(str(filepath), pagesize=A4)
            
            # Prepare content
            story = []
            styles = getSampleStyleSheet()
            
            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            story.append(Paragraph(report_data["title"], title_style))
            story.append(Spacer(1, 20))
            
            # Add metadata
            self._add_pdf_metadata(story, styles, report_data)
            
            # Add executive summary
            self._add_pdf_executive_summary(story, styles, report_data)
            
            # Add execution details
            self._add_pdf_execution_details(story, styles, report_data)
            
            # Add AI analysis
            self._add_pdf_ai_analysis(story, styles, report_data)
            
            # Add performance metrics
            self._add_pdf_performance_metrics(story, styles, report_data)
            
            # Add Playwright code
            self._add_pdf_playwright_code(story, styles, report_data)
            
            # Build PDF
            doc.build(story)
            
            self.logger.info(f"PDF report generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Failed to generate PDF report: {e}", exc_info=True)
            raise
    
    def _add_pdf_metadata(self, story: List, styles: Any, report_data: Dict[str, Any]):
        """Add metadata to PDF report."""
        story.append(Paragraph("Report Information", styles['Heading2']))
        
        # Create metadata table
        data = [
            ["Property", "Value"],
            ["Instructions", report_data["instructions"]],
            ["URL", report_data["url"]],
            ["Status", report_data["status"]],
            ["Timestamp", report_data["timestamp"]],
            ["Execution Time", f"{report_data['execution_time']:.2f} seconds"],
            ["Success Rate", f"{report_data['success_rate']:.1f}%"]
        ]
        
        table = Table(data, colWidths=[2*inch, 4*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
    
    def _add_pdf_executive_summary(self, story: List, styles: Any, report_data: Dict[str, Any]):
        """Add executive summary to PDF report."""
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        
        summary_text = f"""
        <b>Automation Status:</b> {report_data["status"].upper()}<br/><br/>
        """
        
        if report_data["status"] == "completed":
            summary_text += f"""
            ✅ The automation completed successfully with {report_data['success_rate']:.1f}% success rate 
            in {report_data['execution_time']:.2f} seconds.
            """
        else:
            summary_text += f"""
            ❌ The automation encountered issues during execution.<br/>
            Error: {report_data['error_details']}
            """
        
        # Add AI insights
        insights = report_data["ai_insights"]
        if insights["elements_analyzed"] > 0:
            summary_text += f"""
            <br/><br/><b>🤖 AI Analysis:</b> Analyzed {insights['elements_analyzed']} elements 
            with {insights['confidence_score']:.1f}% confidence.
            """
        
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 20))
    
    def _add_pdf_execution_details(self, story: List, styles: Any, report_data: Dict[str, Any]):
        """Add execution details to PDF report."""
        story.append(Paragraph("Execution Details", styles['Heading2']))
        
        if not report_data["steps"]:
            story.append(Paragraph("No execution steps available.", styles['Normal']))
            story.append(Spacer(1, 20))
            return
        
        # Create steps table
        headers = ["Step", "Action", "Description", "Status"]
        data = [headers]
        
        for i, step in enumerate(report_data["steps"], 1):
            data.append([
                str(i),
                step.get("action", ""),
                step.get("description", ""),
                step.get("status", "")
            ])
        
        table = Table(data, colWidths=[0.5*inch, 1*inch, 3*inch, 1*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
    
    def _add_pdf_ai_analysis(self, story: List, styles: Any, report_data: Dict[str, Any]):
        """Add AI analysis to PDF report."""
        story.append(Paragraph("AI Analysis", styles['Heading2']))
        
        insights = report_data["ai_insights"]
        
        analysis_text = f"""
        <b>Element Analysis:</b><br/>
        • Login Elements: {insights['login_elements_found']}<br/>
        • Form Elements: {insights['form_elements_found']}<br/>
        • Navigation Elements: {insights['navigation_elements_found']}<br/>
        • Interactive Elements: {insights['interactive_elements_found']}<br/>
        • Total Elements Analyzed: {insights['elements_analyzed']}<br/>
        • AI Confidence Score: {insights['confidence_score']:.1f}%<br/>
        """
        
        if insights["recommendations"]:
            analysis_text += "<br/><b>AI Recommendations:</b><br/>"
            for rec in insights["recommendations"]:
                analysis_text += f"• {rec}<br/>"
        
        story.append(Paragraph(analysis_text, styles['Normal']))
        story.append(Spacer(1, 20))
    
    def _add_pdf_performance_metrics(self, story: List, styles: Any, report_data: Dict[str, Any]):
        """Add performance metrics to PDF report."""
        story.append(Paragraph("Performance Metrics", styles['Heading2']))
        
        metrics = report_data["performance_metrics"]
        
        # Create metrics table
        data = [
            ["Metric", "Value"],
            ["Total Steps", str(metrics["total_steps"])],
            ["Successful Steps", str(metrics["successful_steps"])],
            ["Failed Steps", str(metrics["failed_steps"])],
            ["Execution Time", f"{metrics['execution_time']:.2f} seconds"],
            ["Average Step Time", f"{metrics['average_step_time']:.2f} seconds"],
            ["Fastest Step", f"{metrics['fastest_step']:.2f} seconds"],
            ["Slowest Step", f"{metrics['slowest_step']:.2f} seconds"]
        ]
        
        table = Table(data, colWidths=[2*inch, 2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
    
    def _add_pdf_playwright_code(self, story: List, styles: Any, report_data: Dict[str, Any]):
        """Add Playwright code to PDF report."""
        story.append(Paragraph("Generated Playwright Code", styles['Heading2']))
        
        # Create code style
        code_style = ParagraphStyle(
            'CodeStyle',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=8,
            leftIndent=20,
            rightIndent=20,
            backColor=colors.lightgrey
        )
        
        # Split code into paragraphs
        code_lines = report_data["playwright_code"].split('\n')
        for line in code_lines:
            if line.strip():
                story.append(Paragraph(line, code_style))
            else:
                story.append(Spacer(1, 6))
        
        story.append(Spacer(1, 20))
    
    def get_available_formats(self) -> List[str]:
        """Get list of available report formats based on installed libraries."""
        return [fmt for fmt, available in self.libraries_available.items() if available]
    
    def install_missing_libraries(self) -> Dict[str, bool]:
        """Install missing report generation libraries."""
        missing_libs = []
        
        if not self.libraries_available["docx"]:
            missing_libs.append("python-docx")
        
        if not self.libraries_available["excel"]:
            missing_libs.extend(["pandas", "openpyxl"])
        
        if not self.libraries_available["pdf"]:
            missing_libs.append("reportlab")
        
        if missing_libs:
            self.logger.info(f"Missing libraries for report generation: {missing_libs}")
            self.logger.info("Please install them using: pip install " + " ".join(missing_libs))
        
        return {
            "docx": self.libraries_available["docx"],
            "excel": self.libraries_available["excel"],
            "pdf": self.libraries_available["pdf"]
        }